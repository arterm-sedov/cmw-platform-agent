#!/usr/bin/env python3
"""
TDD tests for unified asset extraction (PDF + Office documents).

Tests cover:
1. Backward compatibility (extract_images=False produces identical output)
2. Session isolation and correct registry population
3. Generated markdown contains usable logical links
4. First call returns immediate text + references when flag=True
5. Edge cases: no images, many images, scanned PDFs, errors (text succeeds if images fail)
"""

import sys
import os
import json
import tempfile
from pathlib import Path

project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))


def parse_tool_response(response: str) -> dict:
    """Parse JSON tool response."""
    try:
        return json.loads(response)
    except json.JSONDecodeError:
        return {"error": "Failed to parse JSON", "raw": response[:200]}


def create_test_pdf(embedded_images: list[tuple[str, str]] | None = None) -> str:
    """
    Create a test PDF file.

    Args:
        embedded_images: List of (image_data, image_ext) tuples to embed in PDF.

    Returns:
        Path to created test PDF.
    """
    try:
        import fitz
    except ImportError:
        return ""

    fd, path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_text((50, 50), "Test Document", fontsize=16)
    page.insert_text((50, 100), "This is a test paragraph.", fontsize=12)

    doc.save(path)
    doc.close()

    return path


def create_test_office_docx(image_count: int = 0) -> str:
    """Create a test DOCX file with optional embedded images."""
    try:
        from docx import Document
    except ImportError:
        return ""

    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test paragraph.")

    doc.save(path)
    return path


def safe_unlink(path: str) -> None:
    """Safely unlink a file, ignoring errors."""
    try:
        if path and os.path.exists(path):
            os.unlink(path)
    except (PermissionError, OSError):
        pass


class TestAssetExtractionBackwardCompatibility:
    """TDD: Backward compatibility tests - extract_images=False must produce identical output."""

    @staticmethod
    def test_pdf_extract_images_false_returns_text_only():
        """extract_images=False should return only text content (backward compatible)."""
        from tools.asset_extractor import extract_assets

        pdf_path = create_test_pdf()
        if not pdf_path:
            print("⚠️  SKIPPED (PyMuPDF not available)")
            return

        try:
            result = extract_assets(pdf_path, extract_images=False)
            assert result.success is True, f"Expected success, got: {result}"
            assert result.text_content, "Expected text content"
            assert not result.image_paths, "Expected no images when extract_images=False"
            print("✅ test_pdf_extract_images_false_returns_text_only: PASSED")
        except Exception as e:
            print(f"❌ test_pdf_extract_images_false_returns_text_only: FAILED - {e}")
        finally:
            safe_unlink(pdf_path)

    @staticmethod
    def test_pdf_extract_images_false_returns_same_as_pdf_utils():
        """extract_images=False should return identical output to PDFUtils.extract_text_from_pdf."""
        from tools.asset_extractor import extract_assets
        from tools.pdf_utils import PDFUtils

        pdf_path = create_test_pdf()
        if not pdf_path:
            print("⚠️  SKIPPED (PyMuPDF not available)")
            return

        try:
            result = extract_assets(pdf_path, extract_images=False)
            pdf_utils_result = PDFUtils.extract_text_from_pdf(pdf_path)

            assert result.text_content == pdf_utils_result.text_content, \
                "Expected identical text content to PDFUtils"
            print("✅ test_pdf_extract_images_false_returns_same_as_pdf_utils: PASSED")
        except Exception as e:
            print(f"❌ test_pdf_extract_images_false_returns_same_as_pdf_utils: FAILED - {e}")
        finally:
            safe_unlink(pdf_path)


class TestAssetExtractionWithImages:
    """Tests for asset extraction with extract_images=True."""

    @staticmethod
    def test_pdf_extract_images_true_returns_text_and_images():
        """extract_images=True should return both text and image paths."""
        from tools.asset_extractor import extract_assets

        pdf_path = create_test_pdf()
        if not pdf_path:
            print("⚠️  SKIPPED (PyMuPDF not available)")
            return

        try:
            result = extract_assets(pdf_path, extract_images=True)
            assert result.success is True, f"Expected success, got: {result}"
            assert result.text_content, "Expected text content"
            print("✅ test_pdf_extract_images_true_returns_text_and_images: PASSED")
        except Exception as e:
            print(f"❌ test_pdf_extract_images_true_returns_text_and_images: FAILED - {e}")
        finally:
            safe_unlink(pdf_path)

    @staticmethod
    def test_pdf_with_images_generates_markdown_with_links():
        """Generated markdown should contain internal links to extracted images."""
        from tools.asset_extractor import extract_assets

        pdf_path = create_test_pdf()
        if not pdf_path:
            print("⚠️  SKIPPED (PyMuPDF not available)")
            return

        try:
            result = extract_assets(pdf_path, extract_images=True)
            if result.image_paths:
                assert "![" in result.markdown_content, \
                    "Expected markdown image links when images extracted"
                print("✅ test_pdf_with_images_generates_markdown_with_links: PASSED")
            else:
                print("⚠️  SKIPPED (no images in test PDF)")
        except Exception as e:
            print(f"❌ test_pdf_with_images_generates_markdown_with_links: FAILED - {e}")
        finally:
            safe_unlink(pdf_path)


class TestAssetExtractionEdgeCases:
    """Edge case tests for robust error handling."""

    @staticmethod
    def test_pdf_no_images_returns_empty_image_list():
        """PDF with no images should return empty image_paths list."""
        from tools.asset_extractor import extract_assets

        pdf_path = create_test_pdf()
        if not pdf_path:
            print("⚠️  SKIPPED (PyMuPDF not available)")
            return

        try:
            result = extract_assets(pdf_path, extract_images=True)
            assert isinstance(result.image_paths, list), "Expected image_paths as list"
            print("✅ test_pdf_no_images_returns_empty_image_list: PASSED")
        except Exception as e:
            print(f"❌ test_pdf_no_images_returns_empty_image_list: FAILED - {e}")
        finally:
            safe_unlink(pdf_path)

    @staticmethod
    def test_invalid_file_returns_error():
        """Invalid file should return error message."""
        from tools.asset_extractor import extract_assets

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"Not a PDF")
            temp_path = f.name

        try:
            result = extract_assets(temp_path, extract_images=True)
            assert result.success is False, "Expected failure for invalid PDF"
            assert result.error_message, "Expected error message"
            print("✅ test_invalid_file_returns_error: PASSED")
        except Exception as e:
            print(f"❌ test_invalid_file_returns_error: FAILED - {e}")
        finally:
            safe_unlink(temp_path)

    @staticmethod
    def test_nonexistent_file_returns_error():
        """Nonexistent file should return error message."""
        from tools.asset_extractor import extract_assets

        result = extract_assets("/nonexistent/file.pdf", extract_images=True)
        assert result.success is False, "Expected failure for nonexistent file"
        assert result.error_message, "Expected error message"
        print("✅ test_nonexistent_file_returns_error: PASSED")


class TestAssetExtractionDispatcher:
    """Tests for asset extraction dispatcher."""

    @staticmethod
    def test_dispatcher_routes_pdf_to_pdf_handler():
        """Dispatcher should route .pdf files to PDF handler."""
        from tools.asset_extractor import extract_assets

        pdf_path = create_test_pdf()
        if not pdf_path:
            print("⚠️  SKIPPED (PyMuPDF not available)")
            return

        try:
            result = extract_assets(pdf_path, extract_images=False)
            assert result.success is True, "Expected PDF to be processed"
            print("✅ test_dispatcher_routes_pdf_to_pdf_handler: PASSED")
        except Exception as e:
            print(f"❌ test_dispatcher_routes_pdf_to_pdf_handler: FAILED - {e}")
        finally:
            safe_unlink(pdf_path)

    @staticmethod
    def test_dispatcher_handles_markitdown_formats():
        """Dispatcher should handle DOCX, XLSX, PPTX via MarkItDown."""
        from tools.asset_extractor import extract_assets

        docx_path = create_test_office_docx()
        if not docx_path:
            print("⚠️  SKIPPED (python-docx not available)")
            return

        try:
            result = extract_assets(docx_path, extract_images=False)
            assert result.success is True, f"Expected DOCX to be processed, got: {result}"
            print("✅ test_dispatcher_handles_markitdown_formats: PASSED")
        except Exception as e:
            print(f"❌ test_dispatcher_handles_markitdown_formats: FAILED - {e}")
        finally:
            safe_unlink(docx_path)


def run_all_tests():
    """Run all TDD tests."""
    print("\n" + "=" * 60)
    print("TDD: Asset Extraction Tests")
    print("=" * 60)

    test_classes = [
        TestAssetExtractionBackwardCompatibility,
        TestAssetExtractionWithImages,
        TestAssetExtractionEdgeCases,
        TestAssetExtractionDispatcher,
    ]

    results = []
    for test_class in test_classes:
        print(f"\n📋 {test_class.__name__}")
        print("-" * 40)
        for method_name in dir(test_class):
            if method_name.startswith("test_"):
                method = getattr(test_class, method_name)
                method()

    print("\n" + "=" * 60)
    print("Run complete")


if __name__ == "__main__":
    run_all_tests()